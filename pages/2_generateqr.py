import streamlit as st
import sqlite3
import qrcode
import io
import base64
from PIL import Image
import os
from docx import Document
from docx.shared import Inches

DATABASE = os.path.join(os.path.dirname(__file__), '../samples.db')

def main():

    with st.sidebar:
        st.page_link('main.py', label='Guide')
        st.page_link('pages/1_createsample.py', label='🔵 Create Sample')
        st.page_link('pages/2_generateqr.py', label='🔵 Generate QR Code')
        st.page_link('pages/3_shippingstatus.py', label='🚐 Shipping Status')
        st.page_link('pages/4_insertresults.py', label='🔴 Insert Results')
        st.page_link('pages/5_library.py', label='📚 Library')



def get_request_names():
    conn = sqlite3.connect(DATABASE)
    cursor = conn.cursor()
    cursor.execute("SELECT DISTINCT request_name FROM samples")
    request_names = [row[0] for row in cursor.fetchall()]
    conn.close()
    return request_names

def get_samples_by_request_name(request_name):
    conn = sqlite3.connect(DATABASE)
    cursor = conn.cursor()
    cursor.execute("SELECT eln_number, concentration FROM samples WHERE request_name = ?", (request_name,))
    samples = cursor.fetchall()
    conn.close()
    return samples

def generate_qr_code(data):
    qr = qrcode.QRCode(version=1, box_size=10, border=5)
    qr.add_data(data)
    qr.make(fit=True)
    img = qr.make_image(fill='black', back_color='white')
    return img

# Streamlit app
st.title("QR Code Generator for Samples")

# Fetch unique request names
request_names = get_request_names()

if request_names:
    # Dropdown menu for selecting a request name
    selected_request_name = st.selectbox("Select a project (request name):", request_names)
    
    # Fetch samples for the selected request name
    samples = get_samples_by_request_name(selected_request_name)

    if samples:
            qr_images = []
            qr_size_cm = st.slider("Select QR code size (in cm):", min_value=2.0, max_value=10.0, value=5.0, step=0.5)
            qr_size_inches = qr_size_cm / 2.54  # Convert cm to inches

            for eln_number, concentration in samples:
                qr_content = f"ELN Number: {eln_number}, Concentration: {concentration}"
                
                # Generate QR code
                qr_image = generate_qr_code(qr_content)
                qr_images.append((eln_number, concentration, qr_image))
                
                # Convert QR code image to bytes for Streamlit display
                buffer = io.BytesIO()
                qr_image.save(buffer, format="PNG")
                buffer.seek(0)
                qr_image_bytes = buffer.getvalue()
                
                # Display QR code in a smaller size
                st.image(qr_image_bytes, caption=f"ELN: {eln_number} Concentration: {concentration}", width=200)

            # Button to download all QR codes in a Word file
            # Automatically trigger the download of all QR codes as a Word file
            # Create a Word document
            doc = Document()
            doc.add_heading("QR Codes", level=1)

            for eln_number, concentration, qr_image in qr_images:
                buffer = io.BytesIO()
                qr_image.save(buffer, format="PNG")
                buffer.seek(0)
                doc.add_picture(buffer, width=Inches(qr_size_inches))  # Add QR code image to the document with specified size
                
                # Add the text below the QR code
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(f"ELN: {eln_number} Concentration: {concentration}")
                
                # Calculate font size to match the width of the QR code
                font = run.font
                font_size_points = qr_size_cm * 28.35 / len(run.text)  # Approximate font size in points (1 cm = 28.35 points)
                font.size = Inches(font_size_points / 28.35)  # Convert points to inches

            # Save the Word document to a BytesIO buffer
            word_buffer = io.BytesIO()
            doc.save(word_buffer)
            word_buffer.seek(0)

            # Automatically download the Word file with a dynamic name
            file_name = f"sampleqrcode_{selected_request_name}.docx"
            st.download_button(
                label="Download Word File",
                data=word_buffer,
                file_name=file_name,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

if __name__ == "__main__":
    main()